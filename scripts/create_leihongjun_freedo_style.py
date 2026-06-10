from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("generated/雷洪钧教授访问武汉考比锐特_飞渡科技样式版.docx")

TITLE = "以数字孪生链接产业与课堂，雷洪钧教授到访武汉考比锐特"
SUBTITLE = "聚焦工业数据分析与高校教学教材数字孪生应用合作"
META = "考比锐特（武汉）数智科技有限公司 | 2026年6月"

BLOCKS = [
    (
        "cover",
        "从真实机器人到虚拟产线，从工业数据到教学资源，数字孪生正在成为连接先进制造与高校课堂的新入口。",
    ),
    (
        "body",
        "近日，新能源汽车专家、武汉科技大学硕士生导师雷洪钧教授一行莅临考比锐特（武汉）数智科技有限公司参观指导。围绕公司发展、项目实践、工业数据分析以及高校教学教材数字孪生应用等方向，双方展开了深入交流。",
    ),
    (
        "body",
        "本次交流以公司介绍为起点，以爱普生机器人虚实联动项目为技术展示，以工业领域数据分析为业务方向，最终落到高校教学教材数字孪生应用这一未来合作主体。",
    ),
    ("tag", "01  公司介绍"),
    ("heading", "立足武汉，服务工业数智化转型"),
    (
        "body",
        "武汉考比锐特团队首先向雷洪钧教授一行介绍了公司的发展定位与业务布局。考比锐特（武汉）数智科技有限公司依托宁波考比锐特智能科技有限公司在智能装备、数字化工厂和工业软件方面的技术积累，面向制造企业提供数字孪生、虚拟仿真、产线验证和数据分析等数智化服务。",
    ),
    (
        "body",
        "“考比锐特”源于英文 Cooperate，寓意合作、协作与共创。公司希望通过开放合作连接产业现场、高校科研和技术应用，以数字技术帮助制造业提升效率、质量和管理水平。",
    ),
    (
        "body",
        "武汉高校资源密集，汽车产业基础雄厚，智能制造应用场景丰富。立足武汉、服务华中，考比锐特（武汉）将持续推动产业数字化经验向教学、科研和人才培养场景延伸。",
    ),
    ("tag", "02  项目展示"),
    ("heading", "爱普生机器人虚实联动，让工业场景“同步可见”"),
    (
        "body",
        "在公司情况介绍之后，团队重点展示了爱普生机器人虚实联动项目。该项目以真实工业机器人为对象，通过数字孪生模型与现场设备联动，实现机器人状态、运动轨迹、任务流程和运行数据在虚拟空间中的同步呈现。",
    ),
    (
        "body",
        "虚拟场景不仅能够还原机器人本体、末端执行器和作业单元，还能与真实设备动作保持对应关系。操作者可以在数字空间中观察运行过程、验证动作路径、理解工艺节拍，并对设备调试、教学演示和异常分析形成直观支撑。",
    ),
    (
        "highlight",
        "对于企业客户，虚实联动可用于产线调试、运维监控和方案验证；对于高校教学，它也可以转化为机器人课程、智能制造实训和数字化工厂教学中的核心案例。",
    ),
    ("tag", "03  业务方向"),
    ("heading", "面向工业领域的数据分析"),
    (
        "body",
        "项目展示之后，武汉考比锐特进一步介绍了公司未来重点方向：工业领域的数据分析。",
    ),
    (
        "body",
        "工业数字化的价值，不止于三维展示和流程可视化，更在于通过数据采集、建模分析和智能决策，帮助企业发现问题、优化过程并提升生产效率。",
    ),
    (
        "body",
        "围绕这一方向，公司将重点关注设备运行状态分析、工艺参数优化、质量趋势预测、产线节拍评估、能耗与效率分析等场景，持续探索数据驱动的工业应用方案。",
    ),
    ("tag", "04  专家指导"),
    ("heading", "从技术展示走向场景落地"),
    (
        "body",
        "在听取公司介绍和项目展示后，雷洪钧教授结合制造业数字化发展趋势，对公司后续发展提出了指导性建议。",
    ),
    (
        "highlight",
        "数字孪生不能停留在概念展示层面，关键要围绕真实场景解决实际问题。",
    ),
    (
        "body",
        "针对工业数据分析方向，教授建议公司进一步强化“场景牵引、数据支撑、模型驱动”的技术路线，从设备运行、工艺优化、质量控制、生产节拍和运维管理等具体问题切入，沉淀工业机理、专家经验和数据模型，逐步形成具有行业特色的解决方案能力。",
    ),
    (
        "body",
        "针对爱普生机器人虚实联动项目，教授认为该项目具备较好的展示性和教学转化基础。后续可进一步提升工程深度，把机器人动作、工艺任务、控制逻辑、运行数据和异常处理结合起来，使其既能服务企业调试与运维，也能服务高校课堂演示和实训教学。",
    ),
    ("tag", "05  核心议题"),
    ("heading", "高校教学教材数字孪生应用"),
    (
        "body",
        "交流的后半程，双方把重心转向教育领域：数字孪生技术如何真正走进高校课堂，如何服务教材建设、实验教学和实训评价。",
    ),
    (
        "body",
        "当前高校工科教材仍以文字、图纸和静态图片为主要载体。对于工业机器人、自动化产线、汽车制造、数字化工厂等复杂内容，学生仅凭传统教材很难理解设备结构、工艺流程、控制逻辑和数据变化。",
    ),
    (
        "body",
        "数字孪生技术可以把这些内容转化为三维可视、动态交互、任务驱动的教学资源，让教材从“静态知识”升级为“可操作的学习场景”。",
    ),
    ("subheading", "新形态教材建设"),
    (
        "body",
        "建设“纸质教材 + 数字模型 + 交互案例”的一体化教材资源，在教材章节中配套三维模型、工艺动画、虚拟任务和案例二维码。",
    ),
    ("subheading", "虚拟仿真实训开发"),
    (
        "body",
        "面向智能制造、机器人工程、车辆工程、自动化等专业，开发可用于课堂演示、实验教学和综合实训的数字孪生项目。",
    ),
    ("subheading", "真实项目案例转化"),
    (
        "body",
        "把企业真实项目转化为课堂案例，让学生通过爱普生机器人虚实联动等项目理解工业现场。",
    ),
    ("subheading", "AI 教学辅助探索"),
    (
        "body",
        "结合工业数据分析能力，探索学生操作评价、学习过程记录、故障诊断提示和个性化训练推荐。",
    ),
    (
        "highlight",
        "数字孪生教材不是简单把三维模型放进教材，而是要回答：学生需要理解什么知识？学生需要完成什么任务？教师如何评价学生是否真正掌握？",
    ),
    ("tag", "06  未来合作"),
    ("heading", "从一次展示到一套体系"),
    (
        "body",
        "面向后续合作，武汉考比锐特可围绕高校教学教材数字孪生应用持续推进：联合高校教师团队梳理课程知识点，将企业项目拆解为适合教学的案例单元；建设三维模型库、任务库和案例库；开发数字孪生教学平台；引入工业数据分析能力，让学生理解数据采集、建模、分析和决策的完整链路。",
    ),
    (
        "body",
        "合作的落脚点很清晰：把企业的工程经验转化为高校可持续使用的教学资源，把一次项目展示转化为一套课程体系，把单个虚实联动案例扩展为覆盖教材、实验、实训和评价的数字孪生教育解决方案。",
    ),
    (
        "body",
        "此次雷洪钧教授一行莅临指导，为武汉考比锐特进一步明确了产教融合方向。公司将继续以数字孪生为技术基础，以工业数据分析为业务主线，以高校教学教材应用为重要合作场景，携手高校和专家团队，共同探索智能制造人才培养的新路径。",
    ),
    ("closing", "以数字孪生链接产业与课堂，让真实工业场景成为面向未来的教学资源。"),
]


def set_font(run, size=None, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_para_border_left(paragraph, color="2F73D8", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)


def add_block(doc, role, text):
    if role == "tag":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        set_font(r, 10.5, True, "2F73D8")
        return

    if role == "heading":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(9)
        r = p.add_run(text)
        set_font(r, 16, True, "0B2545")
        return

    if role == "subheading":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        set_para_border_left(p)
        r = p.add_run(text)
        set_font(r, 12, True, "1F4D78")
        return

    if role == "highlight":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.right_indent = Cm(0.15)
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(9)
        p.paragraph_format.line_spacing = 1.18
        set_para_shading(p, "EEF5FF")
        set_para_border_left(p, "2F73D8", "16")
        r = p.add_run(text)
        set_font(r, 11.5, True, "174A8B")
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.18

    if role == "cover":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_shading(p, "EEF5FF")
        r = p.add_run(text)
        set_font(r, 13, True, "0B2545")
    elif role == "closing":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(text)
        set_font(r, 12.5, True, "2F73D8")
    else:
        r = p.add_run(text)
        set_font(r, 11, False, "222222")


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.45)
    section.bottom_margin = Cm(2.45)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run(TITLE)
    set_font(r, 20, True, "0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(9)
    r = subtitle.add_run(SUBTITLE)
    set_font(r, 13, False, "2F73D8")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(18)
    r = meta.add_run(META)
    set_font(r, 10.5, False, "667085")

    for role, text in BLOCKS:
        add_block(doc, role, text)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
