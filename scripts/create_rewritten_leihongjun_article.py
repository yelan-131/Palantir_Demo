from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("generated/雷洪钧院士访问武汉考比锐特_高校教材数字孪生应用改写版.docx")


TITLE = "雷洪钧院士莅临武汉考比锐特指导交流"
SUBTITLE = "共探高校教学教材数字孪生应用与智能制造人才培养新路径"


ARTICLE = [
    (
        "近日，英国皇家学会制造工艺院雷洪钧院士一行莅临考比锐特（武汉）数智科技有限公司参观指导。"
        "双方围绕数字孪生技术在高校教学、教材建设、虚拟仿真实训、汽车产业人才培养以及人工智能赋能教育等方向进行了深入交流，"
        "共同探讨如何把先进制造业真实场景转化为可教学、可训练、可评价的数字化课程资源。"
    ),
    (
        "作为宁波考比锐特智能科技有限公司在华中区域的重要布局，考比锐特（武汉）立足数字孪生应用与数字化工厂解决方案，"
        "持续推进制造场景的虚拟仿真、产线验证、工艺优化和智能运维。此次雷洪钧院士到访，不仅是对公司数智化方向的关注与指导，"
        "也为数字孪生技术进一步走进高校课堂、服务教学改革和教材升级提供了新的思路。"
    ),
    ("从产业现场到高校课堂：数字孪生让制造知识真正“活起来”", "h1"),
    (
        "传统工科教学中，教材往往以文字、图纸和静态图片为主要载体。对于汽车制造、智能装备、工业机器人、自动化产线等复杂场景，"
        "学生很难仅凭平面教材理解设备结构、工艺逻辑、节拍协同和故障机理。数字孪生技术的价值，正在于把抽象知识转化为可观察、"
        "可交互、可推演的三维数字场景，让教材从“读得懂”进一步走向“看得见、练得会、用得上”。"
    ),
    (
        "交流中，双方一致认为，高校教学中的数字孪生应用不应停留在三维展示层面，而应围绕课程目标、专业能力和产业需求进行系统设计。"
        "一套高质量的数字孪生教材资源，既要呈现设备和产线的外观结构，更要还原工艺流程、控制逻辑、数据变化和异常处置过程，"
        "帮助学生在虚拟环境中建立工程认知。"
    ),
    ("高校教材数字孪生应用的四个重点方向", "h1"),
    (
        "第一，建设“纸质教材 + 数字模型 + 交互案例”一体化的新形态教材。"
        "在教材关键章节中嵌入三维模型、工艺动画、仿真任务和案例二维码，学生通过电脑、平板或虚拟仿真平台即可查看设备结构、"
        "拆解核心部件、观察工艺路径，并完成与章节知识点对应的交互练习。这样的教材不再只是知识说明书，而是一个可持续更新的数字化学习入口。"
    ),
    (
        "第二，构建面向专业课程的虚拟仿真实训体系。"
        "围绕智能制造工程、机械设计制造及其自动化、车辆工程、机器人工程、工业工程、自动化等专业，高校可将真实产线、典型装备和关键工艺转化为虚拟实训项目。"
        "学生可以在数字孪生场景中完成产线布局、设备调试、机器人轨迹规划、工艺参数调整、异常诊断等任务，在进入真实车间前完成充分训练。"
    ),
    (
        "第三，把汽车与先进制造真实案例转化为课堂教学资源。"
        "汽车产业是数字孪生应用最具代表性的场景之一，从焊装、涂装、总装到检测、物流和运维，每一个环节都包含大量工程知识。"
        "考比锐特可依托自身在数字化工厂和智能装备领域的经验，将产业项目中的典型流程沉淀为教学案例，使学生在学习专业理论的同时理解企业真实问题。"
    ),
    (
        "第四，探索“AI + 数字孪生 + 教材”的智能学习模式。"
        "在数字孪生平台中引入人工智能能力，可进一步实现学习过程记录、操作步骤评价、故障原因提示、个性化训练推荐和课程数据分析。"
        "教师能够更清楚地掌握学生的工程理解程度，学生也能在反复演练中形成从理论分析到实践决策的完整能力链条。"
    ),
    ("让教材服务能力培养，而不只是知识传递", "h1"),
    (
        "雷洪钧院士长期关注先进制造、智能装备与产业人才培养。此次交流中，他特别强调，制造业数字化转型最终需要大量既懂工程、又懂数据、"
        "还具备系统思维的复合型人才。高校教材和课程体系如果仍停留在传统知识讲授层面，将难以适应新一轮智能制造发展的要求。"
    ),
    (
        "数字孪生教材的核心意义，正是把企业场景、工程方法和实践能力前置到教学过程中。"
        "学生通过虚拟产线理解工艺流程，通过数据面板观察生产状态，通过仿真任务验证方案可行性，通过故障案例学习分析路径。"
        "这种学习方式能够有效弥补高校实训设备昂贵、场景更新慢、真实产线难以开放等现实问题，让更多学生在校内就接触到接近产业一线的训练环境。"
    ),
    ("武汉考比锐特的产教融合实践空间", "h1"),
    (
        "面向高校教学与教材建设，考比锐特（武汉）可围绕“课程、教材、平台、实训、师资”五个层面持续发力。"
        "在课程层面，与高校共同开发数字孪生导论、数字化工厂仿真、智能制造系统集成、汽车产线虚拟调试等课程；"
        "在教材层面，联合教师团队建设配套三维模型库、案例库和任务库；在平台层面，提供可用于课堂演示、实验教学和综合实训的数字孪生教学平台；"
        "在实训层面，围绕产业项目设计任务式训练；在师资层面，通过企业工程师参与教学、教师入企实践等方式，推动教学内容与产业技术同步更新。"
    ),
    (
        "考比锐特这一名称源于英文“Cooperate”，寓意合作、协作与共创。"
        "对于高校教学教材数字孪生应用而言，合作尤为关键：企业提供真实场景和工程经验，高校提供教学体系和人才培养需求，"
        "专家学者提供前沿视野和方法指导。三方协同，才能让数字孪生从技术展示走向教学改革，从单个案例走向系统化教材建设。"
    ),
    ("共创智能制造教育新生态", "h1"),
    (
        "此次雷洪钧院士莅临指导，为考比锐特（武汉）进一步拓展高校教学教材数字孪生应用提供了重要启发。"
        "未来，考比锐特（武汉）将继续以数字孪生为核心抓手，以人工智能为创新引擎，聚焦智能制造、汽车产业和高校工程教育，"
        "推动产业技术、教学资源和人才培养深度融合。"
    ),
    (
        "面向数智化时代，教材不应只是记录知识的文本，更应成为连接课堂、实验室和产业现场的桥梁。"
        "考比锐特（武汉）愿与高校、科研院所及产业伙伴携手，把真实制造场景转化为高质量教学资源，"
        "共同培养面向未来的智能制造人才，为中国制造业高质量发展注入源源不断的新动能。"
    ),
]


def set_cell_border(cell, color="DADCE0", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_font(run, size=None, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text, style=None):
    if style == "h1":
        p = doc.add_paragraph(style="Heading 1")
        r = p.add_run(text)
        set_font(r, 15, True, "1F4D78")
        return p
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_font(r, 11, False, "222222")
    return p


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    for style_name, size, color, before, after in [
        ("Heading 1", 15, "1F4D78", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(TITLE)
    set_font(r, 20, True, "0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(SUBTITLE)
    set_font(r, 13, False, "555555")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("考比锐特（武汉）数智科技有限公司 | 2026年6月")
    set_font(r, 10.5, False, "666666")

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(3)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.18
    r = lead.add_run("导语：")
    set_font(r, 11, True, "1F4D78")
    r = lead.add_run(
        "本稿以雷洪钧院士到访武汉考比锐特为背景，重点阐述数字孪生技术在高校教学、教材建设和智能制造人才培养中的应用价值。"
    )
    set_font(r, 11, False, "222222")

    for item in ARTICLE:
        if isinstance(item, tuple):
            add_paragraph(doc, item[0], item[1])
        else:
            add_paragraph(doc, item)

    doc.add_paragraph(style="Heading 1").add_run("高校数字孪生教材建设建议").font.size = Pt(15)
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    headers = ["建设模块", "教学价值", "可落地方向"]
    widths = [Cm(3.2), Cm(5.1), Cm(7.5)]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        shade_cell(cell, "F2F4F7")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, 10.5, True, "1F4D78")

    rows = [
        ("三维模型库", "把设备结构、产线布局和工艺流程转化为可视化资源。", "配套教材章节建设模型二维码、交互拆装和工艺动画。"),
        ("虚拟仿真实训", "让学生在低风险环境中完成调试、排故和方案验证。", "建设汽车制造、工业机器人、数字化工厂等典型实训项目。"),
        ("产业案例库", "把企业真实问题转化为课堂案例，提升工程理解。", "沉淀产线规划、节拍优化、质量预测、设备维护等案例。"),
        ("AI学习评价", "记录学习过程并辅助教师进行能力诊断。", "形成操作评分、错误提示、个性化训练推荐和课程数据分析。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].width = widths[i]
            set_cell_border(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            set_font(r, 10, i == 0, "222222")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("考比锐特（武汉） | 数字孪生赋能高校教学教材建设")
    set_font(r, 9, False, "777777")

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
