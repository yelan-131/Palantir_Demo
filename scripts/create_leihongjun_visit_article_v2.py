from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("generated/雷洪钧院士访问武汉考比锐特_交流纪要_院士指导增强版.docx")

TITLE = "雷洪钧院士莅临武汉考比锐特指导交流"
SUBTITLE = "聚焦工业数据分析与高校教学教材数字孪生应用合作"

SECTIONS = [
    (
        None,
        [
            "近日，英国皇家学会制造工艺院雷洪钧院士一行莅临考比锐特（武汉）数智科技有限公司参观指导。此次交流围绕公司发展情况、数字孪生项目实践、工业领域数据分析方向以及高校教学教材数字孪生应用等内容展开。双方在充分了解企业技术基础和项目能力的基础上，重点探讨了未来在教育领域开展合作的可行路径。",
            "本次交流既是一次面向产业应用的技术展示，也是一次面向产教融合的深入沟通。武汉考比锐特希望依托自身在虚实联动、数字孪生和工业数据分析方面的积累，与高校、专家团队共同推动智能制造教学资源建设，把真实工业场景转化为可教学、可训练、可评价的新形态教材和实训内容。",
        ],
    ),
    (
        "一、公司介绍：立足武汉，服务工业数智化转型",
        [
            "交流伊始，武汉考比锐特团队首先向雷洪钧院士一行介绍了公司的基本情况、发展定位和业务方向。考比锐特（武汉）数智科技有限公司是宁波考比锐特智能科技有限公司在华中区域的重要布局，依托母公司在智能装备、数字化工厂和工业软件方面的技术积累，面向制造企业提供数字孪生、虚拟仿真、产线验证和数据分析等数智化服务。",
            "“考比锐特”源于英文 Cooperate，寓意合作、协作与共创。这一名称也体现了企业的发展理念：通过开放合作连接产业现场、高校科研和技术应用，以数字技术帮助制造业提升效率、质量和管理水平。",
            "武汉作为高校资源密集、汽车产业基础雄厚、智能制造应用场景丰富的城市，为公司开展工业数智化服务和教育领域合作提供了良好土壤。公司将以武汉为支点，持续服务华中区域制造企业和高校机构，推动产业数字化经验向教学、科研和人才培养场景延伸。",
        ],
    ),
    (
        "二、项目展示：爱普生机器人虚实联动应用",
        [
            "在公司情况介绍之后，团队重点展示了爱普生机器人虚实联动项目。该项目以真实工业机器人为对象，通过数字孪生模型与现场设备联动，实现机器人状态、运动轨迹、任务流程和运行数据在虚拟空间中的同步呈现。",
            "在展示过程中，虚拟场景不仅能够还原机器人本体、末端执行器和作业单元，还能够与真实设备动作保持对应关系。通过这种虚实联动方式，操作者可以在数字空间中观察机器人运行过程、验证动作路径、理解工艺节拍，并对设备调试、教学演示和异常分析形成直观支撑。",
            "该项目体现了武汉考比锐特在“看得见的数字孪生”与“用得上的工业应用”之间的结合能力。对于企业客户而言，虚实联动可以用于产线调试、运维监控和方案验证；对于高校教学而言，它同样可以转化为机器人课程、智能制造实训和数字化工厂教学中的核心案例。",
        ],
    ),
    (
        "三、公司方向：面向工业领域的数据分析",
        [
            "在项目展示基础上，武汉考比锐特进一步介绍了公司的未来方向，即面向工业领域的数据分析。公司认为，工业数字化的价值不仅在于三维展示和流程可视化，更在于通过数据采集、建模分析和智能决策，帮助企业发现问题、优化过程并提升生产效率。",
            "工业现场每天产生大量设备数据、工艺数据、质量数据和生产管理数据。如何把这些数据从“记录结果”转变为“指导决策”，是制造企业数字化转型中的关键问题。武汉考比锐特将围绕设备运行状态分析、工艺参数优化、质量趋势预测、产线节拍评估、能耗与效率分析等方向，持续探索数据驱动的工业应用方案。",
            "雷洪钧院士对这一方向给予了关注，并与团队就人工智能、制造工艺知识和工业数据模型的结合进行了交流。双方认为，未来制造业竞争不仅体现在设备自动化水平，更体现在企业是否能够沉淀数据资产、理解工艺规律，并利用数据持续优化制造系统。",
        ],
    ),
    (
        "四、院士指导：从技术展示走向场景落地",
        [
            "在听取公司介绍和项目展示后，雷洪钧院士结合制造业数字化发展趋势，对武汉考比锐特后续发展提出了指导性建议。他指出，数字孪生和工业数据分析不能停留在概念展示层面，关键要围绕真实场景解决实际问题，形成可复制、可推广、可持续迭代的应用方案。",
            "针对工业领域的数据分析方向，院士建议公司进一步强化“场景牵引、数据支撑、模型驱动”的技术路线。企业应从设备运行、工艺优化、质量控制、生产节拍和运维管理等具体问题切入，沉淀工业机理、专家经验和数据模型，逐步形成具有行业特色的解决方案能力。",
            "针对爱普生机器人虚实联动项目，院士认为该项目具备较好的展示性和教学转化基础。后续可进一步提升其工程深度，把机器人动作、工艺任务、控制逻辑、运行数据和异常处理结合起来，使其既能服务企业调试与运维，也能服务高校课堂演示和实训教学。",
            "在教育领域合作方面，院士特别强调，应把高校教学教材数字孪生应用作为一个系统工程推进。数字孪生教材不是简单地把模型放进教材，而是要把产业案例转化为教学任务，把工程过程转化为学习路径，把操作结果转化为评价依据，从而真正服务学生工程能力培养。",
            "院士还建议，武汉考比锐特可以充分利用武汉高校资源丰富的优势，选择智能制造、机器人工程、车辆工程、自动化等专业方向率先开展试点，与高校教师共同打磨课程内容、教材资源和实训平台，形成可示范的产教融合成果。",
        ],
    ),
    (
        "五、教育领域讨论：高校教学教材数字孪生应用",
        [
            "在完成公司介绍和项目展示后，双方把交流重点转向教育领域，特别是高校教学教材数字孪生应用。双方一致认为，这一方向可以成为未来合作的主体：企业提供真实工业场景和数字孪生能力，高校提供课程体系和人才培养需求，专家团队提供学术视野和方法指导，共同建设面向智能制造教育的新形态教学资源。",
            "当前，高校工科教材仍以文字、图纸和静态图片为主要载体。对于工业机器人、自动化产线、汽车制造、数字化工厂等复杂内容，学生仅凭传统教材很难理解设备结构、工艺流程、控制逻辑和数据变化。数字孪生技术可以把这些内容转化为三维可视、动态交互、任务驱动的教学资源，让教材从“静态知识”升级为“可操作的学习场景”。",
            "围绕高校教学教材建设，双方重点讨论了几个方向：一是建设“纸质教材 + 数字模型 + 交互案例”的新形态教材，在教材章节中配套三维模型、工艺动画、虚拟任务和案例二维码；二是开发面向智能制造、机器人工程、车辆工程、自动化等专业的虚拟仿真实训项目；三是把企业真实项目转化为课堂案例，让学生通过爱普生机器人虚实联动等项目理解工业现场；四是结合工业数据分析能力，探索学生操作评价、学习过程记录、故障诊断提示和个性化训练推荐等 AI 教学辅助功能。",
            "这一方向的核心，不是简单把工业场景做成三维展示，而是围绕教学目标进行系统化重构。数字孪生教材应当回答三个问题：学生需要理解什么知识，学生需要完成什么任务，教师如何评价学生是否真正掌握。只有把模型、数据、任务和评价结合起来，数字孪生才能真正服务高校教学改革。",
        ],
    ),
    (
        "六、未来合作：以教材与实训平台共建为主体",
        [
            "面向后续合作，武汉考比锐特可围绕高校教学教材数字孪生应用持续推进。第一，联合高校教师团队梳理课程知识点，将企业项目拆解为适合教学的案例单元；第二，建设工业机器人、数字化工厂、汽车制造等方向的三维模型库、任务库和案例库；第三，开发可用于课堂演示、实验教学和综合实训的数字孪生教学平台；第四，引入工业数据分析能力，让学生在学习过程中理解数据采集、数据建模和数据决策的完整链路。",
            "未来合作的落脚点，是把企业的工程经验转化为高校可持续使用的教学资源，把一次项目展示转化为一套课程体系，把单个虚实联动案例扩展为覆盖教材、实验、实训和评价的数字孪生教育解决方案。",
            "此次雷洪钧院士莅临指导，为武汉考比锐特进一步明确了产教融合方向。公司将继续以数字孪生为技术基础，以工业数据分析为业务主线，以高校教学教材应用为重要合作场景，携手高校和专家团队，共同探索智能制造人才培养的新路径。",
        ],
    ),
]

TABLE_ROWS = [
    ("公司介绍", "说明武汉考比锐特的发展定位、母公司基础、华中区域布局和数智化服务能力。"),
    ("项目展示", "展示爱普生机器人虚实联动项目，体现数字孪生在机器人调试、演示和教学中的应用价值。"),
    ("公司方向", "明确公司面向工业领域的数据分析，聚焦设备、工艺、质量、节拍和效率优化。"),
    ("院士指导", "建议从真实场景切入，强化场景牵引、数据支撑、模型驱动，并推动项目向教学资源转化。"),
    ("教育讨论", "重点讨论高校教学教材数字孪生应用，作为未来合作的主体方向。"),
    ("合作落点", "共建新形态教材、虚拟仿真实训平台、工业案例库和数据驱动教学评价体系。"),
]


def set_font(run, size=None, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="DADCE0", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_font(r, 11, False, "222222")
    return p


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run(text)
    set_font(r, 15, True, "1F4D78")
    return p


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Microsoft YaHei"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string("1F4D78")
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(TITLE)
    set_font(r, 20, True, "0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(13)
    r = subtitle.add_run(SUBTITLE)
    set_font(r, 13, False, "555555")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(16)
    r = meta.add_run("考比锐特（武汉）数智科技有限公司 | 2026年6月")
    set_font(r, 10.5, False, "666666")

    for heading, paragraphs in SECTIONS:
        if heading:
            add_heading(doc, heading)
        for paragraph in paragraphs:
            add_body_paragraph(doc, paragraph)

    add_heading(doc, "交流重点概览")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    headers = ["交流环节", "重点内容"]
    widths = [Cm(3.2), Cm(12.6)]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        shade_cell(cell, "F2F4F7")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_font(r, 10.5, True, "1F4D78")

    for label, detail in TABLE_ROWS:
        cells = table.add_row().cells
        for i, value in enumerate([label, detail]):
            cells[i].width = widths[i]
            set_cell_border(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(value)
            set_font(r, 10, i == 0, "222222")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("武汉考比锐特 | 工业数据分析与高校数字孪生教学应用")
    set_font(r, 9, False, "777777")

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
