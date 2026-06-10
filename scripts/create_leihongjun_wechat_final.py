from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("generated/雷洪钧教授访问武汉考比锐特_公众号风格定稿.docx")

TITLE = "雷洪钧教授莅临武汉考比锐特指导交流"
SUBTITLE = "聚焦工业数据分析与高校教学教材数字孪生应用合作"
META = "考比锐特（武汉）数智科技有限公司 | 2026年6月"

CONTENT = [
    (
        "lead",
        "当工业现场的机器人手臂在虚拟空间中同步运行，当一本教材不再只是纸上的文字、图纸和公式，数字孪生正在让制造知识变得可看、可练、可评价。",
    ),
    (
        "body",
        "近日，新能源汽车专家、武汉科技大学硕士生导师雷洪钧教授一行莅临考比锐特（武汉）数智科技有限公司参观指导。本次交流从公司介绍开始，经过爱普生机器人虚实联动项目展示，进一步延伸到工业领域数据分析和高校教学教材数字孪生应用等方向。",
    ),
    (
        "body",
        "这既是一次面向产业现场的技术交流，也是一次面向未来教育合作的深入沟通。双方在了解公司技术基础、项目能力和发展方向的基础上，重点探讨了如何把真实工业场景转化为高校可持续使用的教学资源。",
    ),
    ("heading", "01 | 立足武汉，服务工业数智化转型"),
    (
        "body",
        "交流伊始，武汉考比锐特团队首先向雷洪钧教授一行介绍了公司的基本情况、发展定位和业务方向。",
    ),
    (
        "body",
        "考比锐特（武汉）数智科技有限公司是宁波考比锐特智能科技有限公司在华中区域的重要布局，依托母公司在智能装备、数字化工厂和工业软件方面的技术积累，面向制造企业提供数字孪生、虚拟仿真、产线验证和数据分析等数智化服务。",
    ),
    (
        "body",
        "“考比锐特”源于英文 Cooperate，寓意合作、协作与共创。这一名称也体现了企业的发展理念：通过开放合作连接产业现场、高校科研和技术应用，以数字技术帮助制造业提升效率、质量和管理水平。",
    ),
    (
        "body",
        "武汉高校资源密集，汽车产业基础雄厚，智能制造应用场景丰富，为公司开展工业数智化服务和教育领域合作提供了良好土壤。公司将以武汉为支点，持续服务华中区域制造企业和高校机构，推动产业数字化经验向教学、科研和人才培养场景延伸。",
    ),
    ("heading", "02 | 项目展示：爱普生机器人虚实联动"),
    (
        "body",
        "在公司情况介绍之后，团队重点展示了爱普生机器人虚实联动项目。",
    ),
    (
        "body",
        "该项目以真实工业机器人为对象，通过数字孪生模型与现场设备联动，实现机器人状态、运动轨迹、任务流程和运行数据在虚拟空间中的同步呈现。",
    ),
    (
        "body",
        "展示过程中，虚拟场景不仅能够还原机器人本体、末端执行器和作业单元，还能与真实设备动作保持对应关系。操作者可以在数字空间中观察机器人运行过程、验证动作路径、理解工艺节拍，对设备调试、教学演示和异常分析形成直观支撑。",
    ),
    (
        "body",
        "这个项目的意义在于：对于企业客户，虚实联动可以用于产线调试、运维监控和方案验证；对于高校教学，它也可以转化为机器人课程、智能制造实训和数字化工厂教学中的核心案例。",
    ),
    ("heading", "03 | 公司方向：工业领域的数据分析"),
    (
        "body",
        "项目展示之后，武汉考比锐特进一步介绍了公司的未来方向：工业领域的数据分析。",
    ),
    (
        "body",
        "公司认为，工业数字化的价值不仅在于三维展示和流程可视化，更在于通过数据采集、建模分析和智能决策，帮助企业发现问题、优化过程并提升生产效率。",
    ),
    (
        "body",
        "工业现场每天产生大量设备数据、工艺数据、质量数据和生产管理数据。如何把这些数据从“记录结果”转变为“指导决策”，是制造企业数字化转型中的关键问题。",
    ),
    (
        "body",
        "围绕这一方向，武汉考比锐特将重点关注设备运行状态分析、工艺参数优化、质量趋势预测、产线节拍评估、能耗与效率分析等场景，持续探索数据驱动的工业应用方案。",
    ),
    ("heading", "04 | 专家指导：从技术展示走向场景落地"),
    (
        "body",
        "在听取公司介绍和项目展示后，雷洪钧教授结合制造业数字化发展趋势，对公司后续发展提出了指导性建议。",
    ),
    (
        "quote",
        "数字孪生不能停留在概念展示层面，关键要围绕真实场景解决实际问题。",
    ),
    (
        "body",
        "针对工业数据分析方向，教授建议公司进一步强化“场景牵引、数据支撑、模型驱动”的技术路线，从设备运行、工艺优化、质量控制、生产节拍和运维管理等具体问题切入，沉淀工业机理、专家经验和数据模型，逐步形成具有行业特色的解决方案能力。",
    ),
    (
        "body",
        "针对爱普生机器人虚实联动项目，教授认为该项目具备较好的展示性和教学转化基础。后续可进一步提升工程深度，把机器人动作、工艺任务、控制逻辑、运行数据和异常处理结合起来，使其既能服务企业调试与运维，也能服务高校课堂演示和实训教学。",
    ),
    (
        "body",
        "在教育领域合作方面，教授特别强调，高校教学教材数字孪生应用应作为一个系统工程推进。数字孪生教材不是简单地把三维模型放进教材，而是要把产业案例转化为教学任务，把工程过程转化为学习路径，把操作结果转化为评价依据。",
    ),
    ("heading", "05 | 核心议题：高校教学教材数字孪生应用"),
    (
        "body",
        "交流的后半程，双方把重心转向教育领域：数字孪生技术如何真正走进高校课堂，如何服务教材建设、实验教学和实训评价？",
    ),
    (
        "body",
        "当前高校工科教材仍以文字、图纸和静态图片为主要载体。对于工业机器人、自动化产线、汽车制造、数字化工厂等复杂内容，学生仅凭传统教材很难理解设备结构、工艺流程、控制逻辑和数据变化。",
    ),
    (
        "body",
        "数字孪生技术可以把这些内容转化为三维可视、动态交互、任务驱动的教学资源，让教材从“静态知识”升级为“可操作的学习场景”。",
    ),
    ("body", "双方重点讨论了四个方向："),
    ("subheading", "1. 新形态教材建设"),
    (
        "body",
        "建设“纸质教材 + 数字模型 + 交互案例”的一体化教材资源，在教材章节中配套三维模型、工艺动画、虚拟任务和案例二维码，让学生可以边学边看、边看边练。",
    ),
    ("subheading", "2. 虚拟仿真实训开发"),
    (
        "body",
        "面向智能制造、机器人工程、车辆工程、自动化等专业，开发可用于课堂演示、实验教学和综合实训的数字孪生项目。",
    ),
    ("subheading", "3. 真实项目案例转化"),
    (
        "body",
        "把企业真实项目转化为课堂案例，让学生通过爱普生机器人虚实联动等项目理解工业现场，理解设备、工艺、数据和管理之间的关系。",
    ),
    ("subheading", "4. AI 教学辅助探索"),
    (
        "body",
        "结合工业数据分析能力，探索学生操作评价、学习过程记录、故障诊断提示和个性化训练推荐等 AI 教学辅助功能。",
    ),
    (
        "quote",
        "数字孪生教材要回答三个问题：学生需要理解什么知识？学生需要完成什么任务？教师如何评价学生是否真正掌握？",
    ),
    (
        "body",
        "只有把模型、数据、任务和评价结合起来，数字孪生才能真正服务高校教学改革，而不是停留在单一展示或技术演示层面。",
    ),
    ("heading", "06 | 未来合作：从一次展示到一套体系"),
    (
        "body",
        "面向后续合作，武汉考比锐特可围绕高校教学教材数字孪生应用持续推进：",
    ),
    ("bullet", "联合高校教师团队梳理课程知识点，将企业项目拆解为适合教学的案例单元"),
    ("bullet", "建设三维模型库、任务库和案例库，覆盖工业机器人、数字化工厂、汽车制造等方向"),
    ("bullet", "开发数字孪生教学平台，用于课堂演示、实验教学和综合实训"),
    ("bullet", "引入工业数据分析能力，让学生理解数据采集、建模、分析和决策的完整链路"),
    (
        "body",
        "合作的落脚点很清晰：把企业的工程经验转化为高校可持续使用的教学资源，把一次项目展示转化为一套课程体系，把单个虚实联动案例扩展为覆盖教材、实验、实训和评价的数字孪生教育解决方案。",
    ),
    (
        "body",
        "此次雷洪钧教授一行莅临指导，为武汉考比锐特进一步明确了产教融合方向。公司将继续以数字孪生为技术基础，以工业数据分析为业务主线，以高校教学教材应用为重要合作场景，携手高校和专家团队，共同探索智能制造人才培养的新路径。",
    ),
    (
        "closing",
        "从看得见的数字孪生，到用得上的教学资源——这条路，我们正在走。",
    ),
]


def set_font(run, size=None, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, role, text):
    if role == "heading":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(9)
        r = p.add_run(text)
        set_font(r, 15, True, "1F4D78")
        return

    if role == "subheading":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        set_font(r, 12, True, "0B2545")
        return

    if role == "quote":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.right_indent = Cm(0.25)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.18
        r = p.add_run(text)
        set_font(r, 11.5, True, "1F4D78")
        return

    if role == "bullet":
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        set_font(r, 11, False, "222222")
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.18
    if role == "lead":
        r = p.add_run(text)
        set_font(r, 12, True, "0B2545")
    elif role == "closing":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(text)
        set_font(r, 12, True, "1F4D78")
    else:
        r = p.add_run(text)
        set_font(r, 11, False, "222222")


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(TITLE)
    set_font(r, 20, True, "0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(9)
    r = subtitle.add_run(SUBTITLE)
    set_font(r, 13, False, "555555")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(17)
    r = meta.add_run(META)
    set_font(r, 10.5, False, "666666")

    for role, text in CONTENT:
        add_text(doc, role, text)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
